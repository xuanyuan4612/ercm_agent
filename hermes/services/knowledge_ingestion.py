"""
知识入库服务 — 文档上传→解析→分块→向量化→入库流水线

基于 doc/agents/10-rag-shared-agent.md §五 设计。
"""

from __future__ import annotations

import hashlib
import io
import re
from datetime import UTC, datetime
from typing import Any

import httpx
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from hermes.core.config import settings
from hermes.core.logging import get_logger
from hermes.db.models.knowledge import KnowledgeDocument

logger = get_logger(__name__)

SUPPORTED_FORMATS = frozenset({"txt", "md", "json", "docx", "pdf"})


class IngestionResult:
    """入库结果"""

    def __init__(
        self,
        success: bool,
        doc_id: str = "",
        chunks_created: int = 0,
        chunks_skipped: int = 0,
        error: str = "",
    ) -> None:
        self.success = success
        self.doc_id = doc_id
        self.chunks_created = chunks_created
        self.chunks_skipped = chunks_skipped
        self.error = error

    def to_dict(self) -> dict[str, Any]:
        return {
            "success": self.success,
            "doc_id": self.doc_id,
            "chunks_created": self.chunks_created,
            "chunks_skipped": self.chunks_skipped,
            "error": self.error,
        }


class KnowledgeIngestionService:
    """知识入库服务

    使用方式:
        service = KnowledgeIngestionService(db_session, minio_client=app.state.minio, es_client=app.state.es)
        result = await service.ingest_file(
            file_content=b"...",
            filename="policy.docx",
            kb_type="analysis",
            client="group",
            org_id="org-001",
            security_level="internal",
        )
    """

    def __init__(
        self,
        db: AsyncSession,
        minio_client: Any = None,
        es_client: Any = None,
    ) -> None:
        self.db = db
        self._minio = minio_client
        self._es_client = es_client

    def _get_search_adapter(self) -> Any:
        """获取 SearchAdapter（延迟初始化）"""
        from hermes.integrations.search_adapter import SearchAdapter
        return SearchAdapter(self._es_client)

    # ── 公开方法 ─────────────────────────────────────────────────

    async def ingest_file(
        self,
        file_content: bytes,
        filename: str,
        kb_type: str,
        client: str = "group",
        org_id: str = "*",
        security_level: str = "internal",
        metadata: dict | None = None,
    ) -> IngestionResult:
        """处理上传文件并入库"""
        # S1: 格式校验
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in SUPPORTED_FORMATS:
            return IngestionResult(success=False, error=f"不支持的文件格式: .{ext}")

        # S2: 大小校验
        max_bytes = settings.INGESTION_MAX_FILE_SIZE_MB * 1024 * 1024
        if len(file_content) > max_bytes:
            return IngestionResult(
                success=False,
                error=f"文件大小 {len(file_content) / 1024 / 1024:.1f}MB 超过限制 {settings.INGESTION_MAX_FILE_SIZE_MB}MB",
            )

        # S3: 内容解析
        try:
            text = self._parse_content(file_content, ext)
        except Exception as e:
            logger.error("ingestion_parse_failed", filename=filename, error=str(e))
            return IngestionResult(success=False, error=f"文档解析失败: {e}")

        if not text or not text.strip():
            return IngestionResult(success=False, error="文档内容为空")

        # S4: 文本清洗
        text = self._clean_text(text)

        # S5: 语义分块
        title = filename.rsplit(".", 1)[0] if "." in filename else filename
        chunks = self._chunk_text(text)

        # S6: 原始文件写入 MinIO
        minio_bucket = ""
        minio_key = ""
        if self._minio:
            try:
                minio_bucket = settings.MINIO_BUCKET
                minio_key = f"knowledge/{kb_type}/{filename}"
                self._minio.put_object(
                    bucket_name=minio_bucket,
                    object_name=minio_key,
                    data=io.BytesIO(file_content),
                    length=len(file_content),
                )
                logger.info("minio_uploaded", bucket=minio_bucket, key=minio_key, size=len(file_content))
            except Exception as e:
                logger.warning("minio_upload_failed", filename=filename, error=str(e))
                # MinIO 失败不阻塞入库

        # S7: 去重检查 + 向量化 + 写入 PostgreSQL
        search_adapter = self._get_search_adapter()
        doc_uuid = ""
        chunks_created = 0
        chunks_skipped = 0
        total = len(chunks)

        for i, chunk_text in enumerate(chunks):
            content_hash_val = hashlib.sha256(chunk_text.encode()).hexdigest()

            # 检查是否已存在
            exists = await self.db.execute(
                select(KnowledgeDocument.id).where(
                    KnowledgeDocument.content_hash == content_hash_val,
                    KnowledgeDocument.kb_type == kb_type,
                    KnowledgeDocument.is_active == True,
                ).limit(1)
            )
            if exists.scalar_one_or_none():
                chunks_skipped += 1
                logger.info("ingestion_chunk_skipped", filename=filename, chunk=i + 1, reason="duplicate")
                continue

            # Embedding 向量化
            embedding = await self._get_embedding(chunk_text)

            # 构建文档记录
            doc = KnowledgeDocument(
                kb_type=kb_type,
                title=title,
                content=chunk_text,
                content_hash=content_hash_val,
                embedding=embedding,
                source_path=filename,
                chunk_index=i + 1,
                total_chunks=total,
                security_level=security_level,
                client=client,
                org_id=org_id,
                approval_status="approved",
                effective_at=datetime.now(UTC),
                metadata_={
                    "source": "manual_upload",
                    "original_filename": filename,
                    "format": ext,
                    "minio_bucket": minio_bucket,
                    "minio_key": minio_key,
                    "uploaded_at": datetime.now(UTC).isoformat(),
                    **(metadata or {}),
                },
            )
            self.db.add(doc)
            await self.db.flush()

            if not doc_uuid:
                doc_uuid = str(doc.id)
            chunks_created += 1

            # S8: 写入 Elasticsearch 全文索引（异步，失败不阻塞）
            chunk_id = f"{doc.id}:{i + 1}"
            await search_adapter.index_document({
                "doc_id": str(doc.id),
                "chunk_id": chunk_id,
                "kb_type": kb_type,
                "title": title,
                "content": chunk_text,
                "content_snippet": chunk_text[:300],
                "source_path": filename,
                "security_level": security_level,
                "client": client,
                "org_id": org_id,
                "approval_status": "approved",
                "chunk_index": i + 1,
                "total_chunks": total,
                "updated_at": datetime.now(UTC).isoformat(),
            })

        logger.info(
            "ingestion_complete",
            filename=filename,
            chunks_created=chunks_created,
            chunks_skipped=chunks_skipped,
        )

        return IngestionResult(
            success=True,
            doc_id=doc_uuid,
            chunks_created=chunks_created,
            chunks_skipped=chunks_skipped,
        )

    async def ingest_text(
        self,
        title: str,
        content: str,
        kb_type: str,
        client: str = "group",
        org_id: str = "*",
        security_level: str = "internal",
        metadata: dict | None = None,
    ) -> IngestionResult:
        """直接入库纯文本知识条目"""
        content_hash_val = hashlib.sha256(content.encode()).hexdigest()

        exists = await self.db.execute(
            select(KnowledgeDocument.id).where(
                KnowledgeDocument.content_hash == content_hash_val,
                KnowledgeDocument.kb_type == kb_type,
                KnowledgeDocument.is_active == True,
            ).limit(1)
        )
        if exists.scalar_one_or_none():
            return IngestionResult(success=False, chunks_skipped=1, error="内容已存在（hash 重复）")

        # 分块
        chunks = self._chunk_text(content)

        embedding = None
        if chunks:
            embedding = await self._get_embedding(chunks[0])

        search_adapter = self._get_search_adapter()
        total_chunks = max(len(chunks), 1)

        doc = KnowledgeDocument(
            kb_type=kb_type,
            title=title,
            content=content if len(chunks) <= 1 else chunks[0],
            content_hash=content_hash_val,
            embedding=embedding,
            source_path=f"text://{title}",
            chunk_index=1,
            total_chunks=total_chunks,
            security_level=security_level,
            client=client,
            org_id=org_id,
            approval_status="approved",
            effective_at=datetime.now(UTC),
            metadata_={
                "source": "text_input",
                "uploaded_at": datetime.now(UTC).isoformat(),
                **(metadata or {}),
            },
        )
        self.db.add(doc)
        await self.db.flush()

        # ES 索引：第一块
        await search_adapter.index_document({
            "doc_id": str(doc.id),
            "chunk_id": f"{doc.id}:1",
            "kb_type": kb_type,
            "title": title,
            "content": chunks[0] if chunks else content,
            "content_snippet": (chunks[0] if chunks else content)[:300],
            "source_path": f"text://{title}",
            "security_level": security_level,
            "client": client,
            "org_id": org_id,
            "approval_status": "approved",
            "chunk_index": 1,
            "total_chunks": total_chunks,
            "updated_at": datetime.now(UTC).isoformat(),
        })

        # 写入剩余 chunks
        for i, chunk_text in enumerate(chunks[1:], 2):
            chunk_hash = hashlib.sha256(chunk_text.encode()).hexdigest()
            chunk_doc = KnowledgeDocument(
                kb_type=kb_type,
                title=title,
                content=chunk_text,
                content_hash=chunk_hash,
                embedding=None,
                source_path=f"text://{title}",
                chunk_index=i,
                total_chunks=len(chunks),
                security_level=security_level,
                client=client,
                org_id=org_id,
                approval_status="approved",
                effective_at=datetime.now(UTC),
                metadata_=doc.metadata_,
            )
            self.db.add(chunk_doc)
            await self.db.flush()

            # ES 索引：后续块
            await search_adapter.index_document({
                "doc_id": str(chunk_doc.id),
                "chunk_id": f"{chunk_doc.id}:{i}",
                "kb_type": kb_type,
                "title": title,
                "content": chunk_text,
                "content_snippet": chunk_text[:300],
                "source_path": f"text://{title}",
                "security_level": security_level,
                "client": client,
                "org_id": org_id,
                "approval_status": "approved",
                "chunk_index": i,
                "total_chunks": len(chunks),
                "updated_at": datetime.now(UTC).isoformat(),
            })

        await self.db.flush()
        return IngestionResult(success=True, doc_id=str(doc.id), chunks_created=total_chunks)

    # ── 内部方法 ─────────────────────────────────────────────────

    @staticmethod
    def _parse_content(file_content: bytes, ext: str) -> str:
        """根据文件类型解析文本内容"""
        if ext in ("txt", "md"):
            return file_content.decode("utf-8", errors="replace")

        if ext == "json":
            import json
            data = json.loads(file_content.decode("utf-8"))
            if isinstance(data, dict):
                title = data.get("title", "")
                content = data.get("content", "")
                return f"{title}\n\n{content}" if title else content
            if isinstance(data, list):
                return "\n\n".join(
                    f"{item.get('title', '')}\n{item.get('content', '')}"
                    for item in data if isinstance(item, dict)
                )
            return str(data)

        if ext == "docx":
            try:
                from docx import Document
                doc = Document(io.BytesIO(file_content))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError as e:
                raise ImportError("python-docx 未安装，无法解析 docx 文件") from e

        if ext == "pdf":
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(file_content))
                return "\n".join(
                    page.extract_text() or ""
                    for page in reader.pages
                )
            except ImportError as e:
                raise ImportError("PyPDF2 未安装，无法解析 pdf 文件") from e

        raise ValueError(f"未知文件格式: {ext}")

    @staticmethod
    def _clean_text(text: str) -> str:
        """文本清洗"""
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()

    def _chunk_text(self, text: str) -> list[str]:
        """语义分块：按段落边界 + 字符数限制切分"""
        chunk_size = settings.INGESTION_CHUNK_SIZE
        chunk_overlap = settings.INGESTION_CHUNK_OVERLAP

        paragraphs = text.split("\n\n")
        chunks: list[str] = []
        current = ""

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 2 <= chunk_size:
                current = f"{current}\n\n{para}" if current else para
            else:
                if current:
                    chunks.append(current)
                if len(para) > chunk_size:
                    sub_chunks = self._split_long_paragraph(para, chunk_size, chunk_overlap)
                    chunks.extend(sub_chunks)
                    current = ""
                else:
                    current = para

        if current:
            chunks.append(current)

        if not chunks:
            chunks = [text[:chunk_size]]

        return chunks

    @staticmethod
    def _split_long_paragraph(text: str, chunk_size: int, overlap: int) -> list[str]:
        """对超长段落按句子边界切分"""
        sentences = re.split(r'(?<=[。！？.!?])\s*', text)
        chunks: list[str] = []
        current = ""

        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            if len(current) + len(sent) <= chunk_size:
                current = f"{current}{sent}" if current else sent
            else:
                if current:
                    chunks.append(current)
                if len(sent) > chunk_size:
                    for i in range(0, len(sent), chunk_size - overlap):
                        chunks.append(sent[i:i + chunk_size])
                    current = ""
                else:
                    current = sent

        if current:
            chunks.append(current)

        return chunks or [text[:chunk_size]]

    @staticmethod
    async def _get_embedding(text: str) -> list[float] | None:
        """获取文本 Embedding 向量"""
        try:
            api_key = settings.EMBEDDING_API_KEY.get_secret_value()
            if not api_key:
                return None

            async with httpx.AsyncClient(timeout=5.0) as client:
                response = await client.post(
                    f"{settings.EMBEDDING_API_BASE}/embeddings",
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={"model": settings.EMBEDDING_MODEL, "input": text},
                )
                if response.status_code == 200:
                    data = response.json()
                    vec = data["data"][0]["embedding"]
                    if len(vec) == settings.EMBEDDING_DIM:
                        return vec
                return None
        except Exception as e:
            logger.warning("ingestion_embedding_failed", error=str(e))
            return None
